import base64
import io
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


class DocxParser:
    async def parse(self, filepath: str) -> dict:
        doc = Document(filepath)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(cells))

        text = "\n\n".join(paragraphs)
        images = self._extract_images(doc)

        result: dict = {"text": text}
        if images:
            result["images"] = images
            result["extractMethod"] = "hybrid" if text.strip() else "vision"
        else:
            result["extractMethod"] = "text"

        return result

    def _extract_images(self, doc: Document) -> list[dict]:
        images: list[dict] = []
        rels = doc.part.rels

        for rel_id, rel in rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    raw = image_part.blob
                    content_type = image_part.content_type or "image/png"
                    b64 = base64.b64encode(raw).decode("ascii")
                    images.append({
                        "base64": b64,
                        "mediaType": content_type,
                    })
                except Exception:
                    continue

        return images
