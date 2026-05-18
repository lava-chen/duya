import os
import tempfile
import subprocess
from typing import Any, Optional


class DocParser:
    """Parser for .doc files using LibreOffice conversion.

    Converts .doc to .docx via LibreOffice headless, then parses with python-docx.
    """

    def __init__(self, libreoffice_path: Optional[str] = None) -> None:
        self._libreoffice_path = libreoffice_path or self._detect_libreoffice()

    def _detect_libreoffice(self) -> Optional[str]:
        import platform
        import shutil

        system = platform.system()
        if system == "Darwin":
            candidates = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            ]
        elif system == "Windows":
            candidates = [
                "C:/Program Files/LibreOffice/program/soffice.exe",
                "C:/Program Files (x86)/LibreOffice/program/soffice.exe",
            ]
        else:
            candidates = [
                "/usr/bin/soffice",
                "/usr/local/bin/soffice",
            ]

        for candidate in candidates:
            if os.path.isfile(candidate):
                return candidate

        soffice = shutil.which("soffice")
        if soffice:
            return soffice

        return None

    async def parse(self, filepath: str) -> dict:
        if not self._libreoffice_path:
            raise RuntimeError(
                "LibreOffice not found — cannot parse .doc files. "
                "Please install LibreOffice or convert the file to .docx format."
            )

        # Create a temporary directory for conversion
        with tempfile.TemporaryDirectory() as tmpdir:
            # Convert .doc to .docx using LibreOffice headless
            # --headless: run without GUI
            # --convert-to: output format
            # --outdir: output directory
            cmd = [
                self._libreoffice_path,
                "--headless",
                "--convert-to", "docx",
                "--outdir", tmpdir,
                filepath,
            ]

            try:
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if result.returncode != 0:
                    raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            except subprocess.TimeoutExpired:
                raise RuntimeError("LibreOffice conversion timed out (30s)")

            # Find the converted .docx file
            # LibreOffice naming: originalname.docx in output dir
            original_name = os.path.splitext(os.path.basename(filepath))[0]
            converted_path = os.path.join(tmpdir, f"{original_name}.docx")

            if not os.path.isfile(converted_path):
                # Try globbing in case of name mangling
                import glob
                matches = glob.glob(os.path.join(tmpdir, "*.docx"))
                if not matches:
                    raise RuntimeError(f"LibreOffice conversion produced no output: {result.stdout} {result.stderr}")
                converted_path = matches[0]

            # Parse the converted .docx with python-docx
            return self._parse_docx(converted_path)

    def _parse_docx(self, docx_path: str) -> dict:
        from docx import Document

        doc = Document(docx_path)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(cells))

        text = "\n\n".join(paragraphs)
        images = self._extract_images(doc)

        result: dict = {"text": text}
        if images:
            result["images"] = images
            result["extractMethod"] = "hybrid" if text.strip() else "vision"
        else:
            result["extractMethod"] = "text"

        return result

    def _extract_images(self, doc) -> list[dict]:
        import base64
        images: list[dict] = []

        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    raw = image_part.blob
                    content_type = image_part.content_type or "image/png"
                    b64 = base64.b64encode(raw).decode("ascii")
                    images.append({
                        "base64": b64,
                        "mediaType": content_type,
                    })
                except Exception:
                    continue

        return images
